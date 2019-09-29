from docx import Document
from docx.shared import Inches
import pandas as pd

class doc_class(object):
    def __init__(self):
        self.document = Document()

    def write_heading(self,heading_str):
        self.document.add_heading(heading_str, 0)
        pass

    def write_picture(self,pic_address,width=5,height=15):
        self.document.add_picture(pic_address, width=Inches(width),height=Inches(height))
        pass

    def write_dataframe(self,df_org:pd.DataFrame):
        df=df_org.reset_index()
        table = self.document.add_table(rows=1, cols=df.shape[1])
        hdr_cells = table.rows[0].cells
        for columns_num in range(df.shape[1]):
            data = df.columns[columns_num]
            hdr_cells[columns_num].text = data
            pass

        for row_num in range(df.shape[0]):
            row_cells = table.add_row().cells
            for columns_num in range(df.shape[1]):
                data = df.iloc[row_num,columns_num]
                row_cells[columns_num].text = str(data)
                pass
            pass
        pass

    def save_file(self,addresss):
        self.document.save(addresss)
        pass

    def document_generation(self):
        factor_name='apm'
        pic_addr=r'D:\code\factor_module\factor_handle\factor_report_picture'

        df=pd.DataFrame([[1,2],[3,4]],index=list('ab'),columns=list('cd'))
        self.write_dataframe(df)
        # self.write_picture('11.jpg')
        self.document.save(pic_addr+'/'+factor_name+'.docx')
        pass

    pass

d=doc_class()
d.document_generation()


